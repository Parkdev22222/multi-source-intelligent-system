"""Generate the patent application as a .docx file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_korean_font(run, font_name="맑은 고딕", size=11, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=True)
    p.paragraph_format.space_after = Pt(12)


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_korean_font(run, size=15, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_korean_font(run, size=13, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_korean_font(run, size=11.5, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_para(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + 0.75 * level)
    run = p.add_run(text)
    set_korean_font(run, size=11)
    p.paragraph_format.line_spacing = 1.4


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Consolas")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.1
    # 음영 추가
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_table(doc, header, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_korean_font(run, size=10.5, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_korean_font(run, size=10)
    if widths:
        for j, w in enumerate(widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph()


def build():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ───────── 제목 ─────────
    add_title(doc, "특 허 출 원 명 세 서", size=20)
    add_para(doc, "")

    # ───────── 발명의 명칭 ─────────
    add_h1(doc, "【발명의 명칭】")
    add_para(
        doc,
        "지식 그래프 기반 시계열 컨텍스트 주입을 이용한 AI 에이전트 군사 판독보고서 자동 생성 시스템 및 그 방법",
        bold=True,
    )
    add_para(
        doc,
        "(영문: AI Agent-Based Military Intelligence Report Generation System and Method "
        "Using Knowledge-Graph-Augmented Temporal Context Injection)",
    )

    # ───────── 1. 개요 ─────────
    add_h1(doc, "1. 개요")

    add_h2(doc, "(1) 수탁과제 무관 근거 설명")

    add_h3(doc, "1) 제안 발명 내용 요약")
    add_para(
        doc,
        "영상 객체 탐지·페어링 결과를 지식 그래프로 누적하고, 그 그래프에서 추출한 시계열 컨텍스트를 "
        "LLM에 주입하여 표준 IMINT 8개 섹션의 군사 판독보고서를 자동 생성하는 AI 에이전트 시스템.",
    )

    add_h3(doc, "2) 중복성·차별성")
    add_bullet(doc, "중복성: 없음")
    add_bullet(
        doc,
        "차별성: 기존 LLM 보고서 자동 생성 특허는 (i) 텍스트 코퍼스 기반 RAG 또는 "
        "(ii) 단발 프레임 탐지만을 LLM 입력으로 사용함. 본 발명은 영상 탐지의 정형 출력을 "
        "LLM 호출 없이 결정론적으로 지식 그래프에 누적하고, Local + Global 이중 검색으로 산출한 "
        "~500토큰 시계열 컨텍스트를 LLM 프롬프트에 prepend하여 시계열 패턴 인텔리전스 보고로 "
        "격상시키는 점에서 차별됨.",
    )

    add_h2(doc, "(2) 발명 제안 배경")
    add_para(
        doc,
        "위성·드론 영상 인텔리전스 분야는 SAM, YOLO-World 등 zero-shot 텍스트 프롬프트 기반 "
        "객체 탐지 모델이 보급되면서 탐지 자동화 수준이 크게 향상되었다. 그러나 군사 분석관이 "
        "요구하는 최종 산출물은 탐지 결과의 나열이 아니라 표준 IMINT 형식의 정형 판독보고서이며, "
        "이는 LLM에 탐지 결과를 단순 입력하는 방식만으로는 정확도·일관성·도메인 적합성을 확보하기 "
        "어렵다.",
    )
    add_para(
        doc,
        "또한 단일 프레임의 탐지 결과만 LLM에 제공하는 경우 '이 지역에서 5회째 반복되는 배치 패턴', "
        "'직전 3회 모두 소실 후 재출현' 같은 시계열 인텔리전스 패턴이 보고서에 반영되지 못해 분석관의 "
        "후속 판단을 보조하지 못한다. 본 발명은 (i) 영상 탐지·페어링 결과를 LLM 호출 없이 그래프로 "
        "결정론적 누적하고, (ii) 그래프 기반 지역·전역 이중 검색으로 시계열 컨텍스트를 추출하며, "
        "(iii) 변화 객체(신규·소실)만을 표준 8섹션 프롬프트와 결합하여 LLM에게 전달하는 AI 에이전트 "
        "시스템을 제안한다.",
    )

    add_h2(doc, "(3) 산업상 이용 분야 (용도)")
    for s in [
        "국방·정보 분야 위성/드론 영상 판독 자동화 시스템",
        "국가정보·해양감시·국경감시 등 ISR(Intelligence, Surveillance, Reconnaissance) 자동 보고 시스템",
        "재난·산림·해상사고 모니터링용 정형 보고서 자동 생성 시스템",
        "민간 위성영상 서비스(농업·도시 변화·공급망)의 시계열 변화 자동 보고서",
        "자율감시 로봇·드론 군집의 임무 결과 보고서 자동화",
    ]:
        add_bullet(doc, s)

    add_h2(doc, "(4) 종래 기술")

    add_h3(doc, "US 12,488,225 B1 — Modular Open System Architecture for Common Intelligence Picture Generation (Booz Allen Hamilton, 2025년 등록)")
    add_para(
        doc,
        "본 발명과 가장 직접적으로 비교 가능한 실제 등록 특허는 미국 특허 US 12,488,225 B1 "
        "(이하 '인용발명')이다. 인용발명은 2025년에 등록된 비교적 최근의 정부 정보(인텔리전스) 분야 "
        "특허로서, 다중 출처(multi-source) 정보 데이터를 통합하여 공통 인텔리전스 픽처(Common "
        "Intelligence Picture, CIP)를 자동 생성하는 모듈형 개방 시스템 아키텍처(MOSA)에 관한 것이며, "
        "그 핵심 구성에 대규모 언어 모델(LLM)을 다중 정보 융합 엔진으로 활용한다. 이는 본 발명이 "
        "속하는 \"위성영상 → LLM 에이전트 변화탐지 → 객체 이력 기반 판독보고서 자동 생성\" "
        "카테고리에서 가장 가까운 실제 등록 특허에 해당한다.",
    )
    add_bullet(
        doc,
        "[인용특허] US Patent 12,488,225 B1, \"Modular open system architecture for common "
        "intelligence picture generation\", 양수인: Booz Allen Hamilton Inc. (추정), "
        "USPTO 등록 2025.",
    )
    add_bullet(
        doc,
        "[비특허문헌 1] Booz Allen Hamilton, \"Booz Allen i2S2: Interactive AI for Space\", "
        "공식 제품 페이지, boozallen.com/markets/space/i2s2-interactive-ai-for-space.html — "
        "인용특허의 상용 구현체에 해당.",
    )
    add_bullet(
        doc,
        "[비특허문헌 2] Booz Allen Hamilton, \"Booz Allen Deploys the Power of Generative AI "
        "in Space\", investors.boozallen.com 보도자료.",
    )
    add_bullet(
        doc,
        "[비특허문헌 3] Booz Allen Hamilton & Meta, \"AI Vision Language Model for Space\" "
        "공동 시연 발표, BusinessWire 2025.04.25.",
    )
    add_para(
        doc,
        "(출원 단계에서 변리사가 USPTO 및 KIPRIS에서 정확한 양수인·발명자·청구항 전문을 "
        "재확인 권장.)",
    )

    add_para(doc, "인용발명(US 12,488,225)의 핵심 구성은 다음과 같다.")
    add_bullet(
        doc,
        "(가) 모듈형 개방 시스템 아키텍처(MOSA)를 기반으로 위성·OSINT·CYBER·SIGINT 등 다중 출처 "
        "정보를 단일 공통 인텔리전스 픽처(CIP)로 통합한다.",
    )
    add_bullet(
        doc,
        "(나) 다중 정보 융합 LLM(multi-source intelligence fusion LLM)이 우선 정보 요구사항"
        "(Priority Intelligence Requirements, PIR)을 자연어로 입력받아 대화형 질의응답을 통해 "
        "분석관과 상호작용한다.",
    )
    add_bullet(
        doc,
        "(다) LLM 에이전트가 PIR을 분석하여 위성 임무 부여(satellite tasking), OSINT 수집, "
        "CYBER 작전, 기타 정보수집 활동의 적절한 대응 방안(courses of action, COA)을 자동 생성한다.",
    )
    add_bullet(
        doc,
        "(라) 위성 데이터 획득 최적화 플랫폼이 기상·궤도·센서 능력을 평가하여 다중 정보 출처에 대한 "
        "수집 실현 가능성을 산출한다.",
    )
    add_bullet(
        doc,
        "(마) 컨테이너화된 분석 워크벤치(containerized analytics workbench)와 연동하여 "
        "맞춤형 정보 산출물(PDF, 프레젠테이션 등) 형태의 자동 보고서를 생성한다.",
    )

    add_h2(doc, "(5) 종래 기술의 문제점 / 한계")
    add_para(
        doc,
        "인용발명(US 12,488,225)은 본 발명과 동일한 \"LLM + 다중 출처 정보 → 자동 보고서 생성\" "
        "카테고리에 속하는 가장 가까운 등록 특허이나, 본 발명이 해결하는 \"위성영상 객체 단위 시계열 "
        "변화탐지 및 그래프 기반 객체 이력 누적 보고\"라는 좁고 구체적인 도메인에서 비교할 때 "
        "아래와 같은 7가지 본질적 한계를 갖는다.",
    )
    problems = [
        ("정보 융합 수준 — 다중 출처 추상 융합 vs 객체 단위 시계열 변화",
         "인용발명은 위성·OSINT·CYBER·SIGINT 등 이질적 정보 출처를 통합 공통 인텔리전스 픽처(CIP)로 "
         "추상 융합하는 데 초점이 있으며, 영상 내 개별 객체(전차·항공기·시설 등) 단위로 시계열 변화를 "
         "추적하는 구성이 청구범위 및 명세서에 정의되어 있지 않다. 본 발명은 SAM3 등 텍스트 프롬프트 "
         "탐지기로 영상 내 객체를 분할·식별한 후, 객체별 페어링(new/matched/moved/disappeared)을 "
         "수행하여 객체 단위 변화 의미를 LLM에 전달한다."),
        ("입력 패러다임 — PIR 기반 질의응답 vs 자동 변화탐지 트리거",
         "인용발명은 분석관이 자연어로 입력한 우선 정보 요구사항(PIR)에 응답해 LLM 에이전트가 "
         "정보수집·분석을 수행하는 질의응답형 구조이다. 본 발명은 신규 위성영상이 수집되면 자동으로 "
         "탐지·페어링·그래프 갱신·보고서 생성이 트리거되는 사건 기반(event-driven) 파이프라인으로, "
         "분석관의 수동 질의 없이도 변화 보고서를 산출한다."),
        ("객체 이력 저장 — CIP(평탄 통합 픽처) vs 그래프 카운터 누적 메모리",
         "인용발명의 공통 인텔리전스 픽처(CIP)는 다중 출처 정보를 통합 표시하는 \"현재 시점의 "
         "스냅샷\" 구조이며, 동일 객체의 반복 관측을 객체별 누적 카운터(new_count, "
         "disappeared_count 등)로 압축 저장하는 결정론적 그래프 인덱싱 구성이 부재하다. "
         "본 발명은 (자산 노드 × 위치 노드)에 시계열 카운터를 누적하여 \"이 지역 N번째 출현, "
         "K번째 소실 후 재등장\" 같은 시계열 통계를 단일 노드 O(1) 조회로 산출한다."),
        ("자산 간 공출현 패턴 자동 발견 부재 (Louvain)",
         "인용발명은 LLM이 PIR을 분석해 대응 방안(COA)을 생성하나, 관측 데이터 자체에서 자산 "
         "클래스 간 공출현 패턴(기갑+APC+포병 = 기갑 복합체, 레이더+지휘소 = C2 인프라 등)을 그래프 "
         "알고리즘으로 자동 군집화·발견하는 구성이 정의되어 있지 않다. 본 발명은 co_occurred_with "
         "엣지 가중치에 Louvain 알고리즘을 적용해 doctrine 패턴을 자동 식별하고, 분석관 질의 없이도 "
         "패턴을 보고서 컨텍스트로 LLM에게 제시한다."),
        ("도메인 안전성 가드레일 부재 — 'DISAPPEARED' 의미 제약",
         "인용발명은 범용 정보 융합 LLM으로, 영상 변화탐지 특유의 의미론적 제약 — \"영상에서 객체가 "
         "더 이상 관측되지 않음(DISAPPEARED)\"을 \"파괴 확인(destroyed)\"으로 표현해서는 안 됨 — "
         "을 시스템 프롬프트 차원에서 강제하는 구성이 정의되어 있지 않다. 본 발명은 이 의미 제약을 "
         "LLM 시스템 프롬프트에 명시 강제하여 미관측을 파괴로 오해석하는 군사적 오판을 차단한다."),
        ("변화 객체 한정 전달 + 사전 압축 컨텍스트 블록 부재",
         "인용발명은 LLM이 PIR 범위 내 모든 정보를 조회·융합하므로 LLM 호출 비용·지연·응답 가변성이 "
         "크다. 본 발명은 (i) 페어링 결과 중 활동 지표인 new·disappeared만 추출하고 (ii) Local + "
         "Global 그래프 검색으로 ~500 토큰의 압축된 결정론적 historical context 블록을 사전 산출하여 "
         "단일 LLM 호출로 보고서를 생성하므로 토큰·지연·재현성에서 모두 우위이다."),
        ("출력 형식 강제 — 맞춤형 PDF/PPT vs 표준 IMINT 8섹션 + 동일 LLM 번역",
         "인용발명의 보고서는 분석관이 지정한 맞춤형 PDF/프레젠테이션 형식이다. 본 발명은 "
         "CLASSIFICATION / EXECUTIVE SUMMARY / SITUATION / CHANGE ANALYSIS / THREAT ASSESSMENT / "
         "INTELLIGENCE GAPS / RECOMMENDED ACTIONS / APPENDIX의 8개 표준 IMINT 섹션을 시스템 프롬프트 "
         "차원에서 강제하고, 동일 LLM의 재호출을 통한 한국어 번역(섹션 헤더·좌표·신뢰도 토큰 보존)으로 "
         "다국어 분석관 표준 양식 일관성을 확보한다."),
    ]
    add_table(doc, ["#", "인용발명 (US 12,488,225)의 한계", "본 발명의 대응 구성"],
              [[str(i + 1), t, d] for i, (t, d) in enumerate(problems)],
              widths=[0.8, 4.8, 10.4])

    add_para(
        doc,
        "따라서 인용발명 US 12,488,225 B1은 본 발명과 동일한 \"LLM 기반 다중 출처 정보 → 자동 "
        "인텔리전스 보고서\" 카테고리에서 가장 가까운 실제 등록 특허이나, 본 발명이 해결하는 "
        "위성영상 객체 시계열 변화탐지 도메인의 (i) 객체 단위 페어링 기반 변화 의미 추출, "
        "(ii) 영상 수집 사건 기반 자동 트리거, (iii) 그래프 카운터 기반 시계열 누적 메모리, "
        "(iv) Louvain 기반 자산 공출현 패턴 자동 발견, (v) 'DISAPPEARED ≠ destroyed' 도메인 "
        "의미 가드레일, (vi) new/disappeared 한정 + 사전 압축 historical context 블록, "
        "(vii) IMINT 8섹션 강제 및 동일 LLM 번역의 7가지 차별 구성을 모두 결여하므로 본 발명의 "
        "진보성이 인정된다.",
    )

    # ───────── 2. 상세 설명 ─────────
    add_h1(doc, "2. 상세 설명")

    add_h2(doc, "(1) 발명의 내용 (구성)")
    add_para(doc, "본 발명은 다음 5개 기능 모듈을 포함한다.")

    add_h3(doc, "A. 결정론적 그래프 인덱서 (Graph Indexer) [200]")
    for s in [
        "위경도를 0.01°(약 1km) 격자로 양자화하여 위치(location) 노드 키 생성",
        "(객체 클래스 × 위치) 고유 쌍으로 자산(asset) 노드 키 생성",
        "status 필드(new/matched/moved/disappeared)를 자산 노드 카운터 속성에 LLM 호출 없이 결정론적 누적",
        "asset → location 의 found_at 엣지, asset ↔ asset 의 co_occurred_with 엣지를 누적 가중치로 갱신",
    ]:
        add_bullet(doc, s)

    add_h3(doc, "B. 커뮤니티 탐지부 (Community Detector) [300]")
    for s in [
        "co_occurred_with 엣지 가중치 기반 Louvain 알고리즘으로 자산 군집 자동 발견",
        "각 군집의 member_summary(자산 분포·위치 클러스터·관측/신규/소실 통계)를 LLM 비호출로 구조 생성",
    ]:
        add_bullet(doc, s)

    add_h3(doc, "C. 이중 모드 그래프 검색부 (Local + Global Retriever) [500]")
    for s in [
        "Local Search: 현재 좌표 반경 R 내 자산 노드 이력 통계 조회",
        "Global Search: 동일 반경과 중첩되는 커뮤니티의 member_summary 조회",
        "두 결과를 토큰 예산(~500 tokens) 이내로 절단·머지하여 단일 컨텍스트 블록 생성",
    ]:
        add_bullet(doc, s)

    add_h3(doc, "D. 변화 객체 중심 프롬프트 조립부 (Prompt Composer) [600]")
    for s in [
        "페어링 결과 중 'new'·'disappeared' 상태만 추출, 신뢰도 내림차순 정렬, 상위 20건 + 잔여 요약",
        "시스템 프롬프트에 군사 의미 제약 강제: \"'DISAPPEARED'는 미관측을 의미하며 파괴 확인이 아님\"",
        "표준 IMINT 8개 섹션 강제: CLASSIFICATION / EXECUTIVE SUMMARY / SITUATION / CHANGE ANALYSIS / "
        "THREAT ASSESSMENT / INTELLIGENCE GAPS / RECOMMENDED ACTIONS / APPENDIX",
        "C 모듈의 컨텍스트 블록을 사용자 프롬프트 선두에 prepend",
    ]:
        add_bullet(doc, s)

    add_h3(doc, "E. 동일 LLM 재호출 기반 번역·헤더 부착부 [700]")
    for s in [
        "동일 LLM 모델을 1차(영문 보고서) → 2차(한국어 번역)로 재호출",
        "번역 규칙: 섹션 헤더·좌표·타임스탬프·클래스명·신뢰도 토큰 보존, 서술 텍스트만 번역",
        "메타 헤더(모델명·세션ID·관측 시각·페어링 통계) 부착하여 최종 산출",
    ]:
        add_bullet(doc, s)

    add_h2(doc, "(2) 발명의 효과")
    effects = [
        ("시계열 인텔리전스 격상",
         "단발 프레임 변화 보고가 아닌 'N번째 반복 출현, K번 소실 후 재등장' 수준의 시계열 패턴 보고 자동 생성."),
        ("결정론·재현성",
         "그래프 인덱싱이 LLM 비호출 결정론 방식 → 동일 입력에 동일 그래프, 군사 감사 가능성 확보."),
        ("운용 비용 절감",
         "임베딩 모델·벡터 DB·임베딩 GPU 추론 불필요. SQLite + 그래프 알고리즘만으로 RAG 컨텍스트 산출."),
        ("토큰 효율 극대",
         "정지·이동 객체 제외 + ~500 토큰 그래프 컨텍스트만 추가 → 보고서 깊이 강화 + 토큰 절감."),
        ("군사 오판 방지",
         "'DISAPPEARED ≠ destroyed' 의미 제약으로 LLM 군사적 오해석 위험 억제."),
        ("다국어 일관성",
         "동일 LLM 재호출 번역 + 토큰 보존 규칙으로 별도 번역 모델 없이 형식 일관성 확보."),
        ("그래프 패턴 자동 발견",
         "Louvain 군집화로 '기갑·포병 복합체', 'C2 인프라' 같은 doctrine 패턴 자동 식별."),
    ]
    add_table(doc, ["#", "효과", "내용"],
              [[str(i + 1), t, d] for i, (t, d) in enumerate(effects)],
              widths=[1.0, 4.0, 11.0])

    # ───────── 3. 도면 ─────────
    add_h1(doc, "3. 도면")

    add_h2(doc, "(1) 개략적인 전체 그림 — 도 1")
    add_code_block(
        doc,
        """┌──────────────────────────────────────────────────────────────────────────┐
│        AI 에이전트 기반 판독보고서 자동 생성 시스템 [1000]                │
│                                                                          │
│  [입력 측 — 종래기술 활용]                                                │
│  영상 객체 탐지부 [110] → 시계열 페어링부 [120]                          │
│                  │                                                       │
│                  ▼  PairingRecord 배치                                   │
│  [본 발명의 핵심부]                                                       │
│  A. 결정론적 그래프 인덱서 [200]                                          │
│     ─ 격자 양자화부 [210]                                                │
│     ─ 자산 노드 누적부 [220]   ─ 엣지 누적부 [230]                       │
│                  │                                                       │
│                  ▼                                                       │
│  B. 커뮤니티 탐지부 [300]                                                │
│     ─ Louvain 군집화 [310]    ─ member_summary 생성부 [320]              │
│                  │                                                       │
│                  ▼                                                       │
│  지식 그래프 저장소 [400]                                                │
│  (graph_entities / graph_relations / graph_communities)                  │
│                  │                                                       │
│                  ▼                                                       │
│  C. 그래프 검색부 [500]                                                  │
│     ─ Local Search [510]      ─ Global Search [520]                      │
│     ─ 토큰 예산 컨텍스트 생성부 [530]                                     │
│                  │                                                       │
│                  ▼                                                       │
│  D. 변화 객체 중심 프롬프트 조립부 [600]                                 │
│     ─ new/disappeared 추출 [610]  ─ 시스템 프롬프트 제약 [620]            │
│     ─ 8섹션 강제부 [630]          ─ 컨텍스트 prepend [640]                │
│                  │                                                       │
│                  ▼                                                       │
│  E. AI 에이전트 LLM 호출부 [700]                                          │
│     ─ 영문 보고서 생성 [710]    ─ 한국어 번역 [720]                      │
│     ─ 메타 헤더 조립부 [730]                                             │
│                  │                                                       │
│                  ▼                                                       │
│        최종 IMINT 판독보고서 [800] → 보고서 DB                            │
└──────────────────────────────────────────────────────────────────────────┘""",
    )

    add_h2(doc, "(2) 발명의 특징부 세부 도면")

    add_h3(doc, "도 2 — A. 결정론적 그래프 인덱서 [200] 내부 동작")
    add_code_block(
        doc,
        """PairingRecord 입력
(current_class, current_lat/lon, past_class, past_lat/lon,
 status, confidence, session_id)
        │
        ▼
┌─────────────────────────────────────┐
│ [210] 격자 양자화부                  │
│  loc_key = "loc:{lat:.2f},{lon:.2f}" │
│  (소수점 2자리 = 약 1km 격자)         │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [220] 자산 노드 upsert (LLM 비호출)  │
│  asset_id = "asset:{class}:{loc_key}"│
│                                     │
│  status에 따라 카운터 증가:           │
│    new → new_count++                │
│    matched → matched_count++        │
│    moved → moved_count++            │
│    disappeared → disappeared_count++│
│                                     │
│  total_confidence += conf           │
│  sessions.append(session_id)        │
│  first/last_seen 갱신               │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [230] 엣지 누적부                    │
│  found_at: asset → location          │
│            properties.count++         │
│  co_occurred_with: asset ↔ asset      │
│   (같은 session_id + 같은 loc_key)    │
│   properties.count++                  │
└─────────────────┬───────────────────┘
                  ▼
          지식 그래프 저장소 [400]""",
    )

    add_h3(doc, "도 3 — C. 그래프 검색부 [500] 내부 동작")
    add_code_block(
        doc,
        """현재 페어링 중심 좌표 (lat_c, lon_c)
        │
        ▼
┌─────────────────────────────────────┐
│ [510] Local Search                   │
│   - 반경 R(예: 0.05°) 내 asset 노드   │
│   - 자산별 누적 통계 조회             │
│     (new/matched/moved/disappeared,  │
│      avg_conf, first/last_seen)      │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [520] Global Search                  │
│   - 반경 R과 중첩되는 community 조회 │
│   - member_summary 추출              │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [530] 컨텍스트 블록 생성부            │
│   - 토큰 예산(예: 500 tokens) 절단    │
│   - "=== GRAPHRAG HISTORICAL         │
│      CONTEXT ===" 포맷화              │
└─────────────────┬───────────────────┘
                  ▼
        D. 프롬프트 조립부 [600] 로 전달""",
    )

    add_h3(doc, "도 4 — D + E. 프롬프트 조립부 및 AI 에이전트 LLM 호출부")
    add_code_block(
        doc,
        """페어링 배치 입력
        │
        ▼
┌─────────────────────────────────────┐
│ [610] 변화 객체 추출                 │
│   - status ∈ {new, disappeared}     │
│   - confidence 내림차순 정렬          │
│   - 상위 20건 + "외 N건" 요약          │
└─────────────────┬───────────────────┘
                  │  +  [500]의 컨텍스트 블록
                  ▼
┌─────────────────────────────────────┐
│ [620] 시스템 프롬프트 제약            │
│  "You are a military IMINT analyst.  │
│   DISAPPEARED ≠ destroyed …"         │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [630] 8섹션 구조 강제부               │
│  1. CLASSIFICATION                   │
│  2. EXECUTIVE SUMMARY                │
│  3. SITUATION                        │
│  4. CHANGE ANALYSIS                  │
│  5. THREAT ASSESSMENT                │
│  6. INTELLIGENCE GAPS                │
│  7. RECOMMENDED ACTIONS              │
│  8. APPENDIX                         │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [640] 컨텍스트 prepend → 최종 prompt │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [710] LLM 1차 호출 → 영문 보고서      │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [720] 동일 LLM 2차 호출 → 한국어 번역 │
│  섹션헤더/좌표/클래스명 보존          │
└─────────────────┬───────────────────┘
                  ▼
┌─────────────────────────────────────┐
│ [730] 메타 헤더 부착                 │
│  (모델명, 세션ID, 관측시각, 통계)    │
└─────────────────┬───────────────────┘
                  ▼
        최종 IMINT 보고서 [800]""",
    )

    add_h2(doc, "(3) 발명의 설명 (동작 순서도) — 도 5")
    add_code_block(
        doc,
        """              ( 시작 )
                 │
                 ▼
  S100 : 페어링 결과 배치 수신
                 │
                 ▼
  S200 : 격자 양자화 → loc_key 생성
                 │
                 ▼
  S300 : 자산 노드 upsert (카운터 +1, LLM 비호출)
                 │
                 ▼
  S400 : found_at / co_occurred_with 엣지 갱신
                 │
                 ▼
  S500 : Louvain 커뮤니티 탐지
                 │
                 ▼
  S600 : Local Search (반경 R 자산 이력)
                 │
                 ▼
  S700 : Global Search (중첩 커뮤니티 요약)
                 │
                 ▼
  S800 : 토큰 예산 컨텍스트 블록 생성
                 │
                 ▼
  S900 : new/disappeared 추출 + 신뢰도 정렬
                 │
                 ▼
  S1000: 시스템 프롬프트 + 8섹션 강제 + 컨텍스트 prepend
                 │
                 ▼
  S1100: LLM 1차 호출 → 영문 IMINT 보고서
                 │
                 ▼
  S1200: 동일 LLM 2차 호출 → 한국어 번역
                 │
                 ▼
  S1300: 메타 헤더 부착 → 보고서 DB 저장
                 │
                 ▼
              ( 종료 )""",
    )

    add_h2(doc, "(4) 도면에 대한 상세 설명")

    add_h3(doc, "도 1 (전체 시스템)")
    add_para(
        doc,
        "본 시스템 [1000]은 입력 측 종래 기술인 영상 객체 탐지부 [110] 및 시계열 페어링부 [120]로부터 "
        "페어링 레코드(PairingRecord)를 입력받아, 본 발명의 핵심 5개 모듈(A~E)을 순차 진행해 최종 "
        "IMINT 판독보고서 [800]을 산출한다. 입력 측 탐지기는 SAM3, YOLO-World 등으로 대체 가능하며 "
        "본 청구범위에 종속되지 않는다. 본 발명의 권리범위는 모듈 [200]~[700]이 형성하는 결합 구조에 있다.",
    )

    add_h3(doc, "도 2 (그래프 인덱서)")
    add_para(
        doc,
        "격자 양자화부 [210]은 GPS 측정 오차 및 미세 좌표 변동을 단일 위치 노드로 흡수하여 동일 자산의 "
        "반복 관측을 누적 카운터로 압축한다. 자산 노드 upsert부 [220]은 LLM 호출 없이 페어링 상태 필드를 "
        "결정론적으로 가산하므로 동일 입력에 매번 동일한 그래프를 산출한다. 엣지 누적부 [230]은 자산-위치 "
        "found_at 관계와 동일 세션·동일 격자 공출현 자산 쌍의 co_occurred_with 관계를 누적하여 Louvain "
        "알고리즘이 가중치 기반 군집화를 수행할 수 있도록 한다.",
    )

    add_h3(doc, "도 3 (그래프 검색부)")
    add_para(
        doc,
        "Local Search [510]은 보고서 대상 좌표 반경 내 자산 노드 이력을 직접 조회하여 \"이 지역에서 "
        "이 자산이 어떤 빈도·패턴으로 출현·소실되었는가\"를 LLM이 참고할 형태로 추출한다. "
        "Global Search [520]은 동일 반경과 중첩되는 군집의 member_summary를 조회하여 \"이 지역이 어느 "
        "doctrine 패턴(예: 기갑·포병 복합체)에 속하는가\"라는 메타 수준 컨텍스트를 제공한다. "
        "컨텍스트 블록 생성부 [530]은 두 결과를 사전 설정 토큰 예산 이내로 압축한다.",
    )

    add_h3(doc, "도 4 (프롬프트 조립 + LLM 호출)")
    add_para(
        doc,
        "변화 객체 추출부 [610]은 활동 지표인 신규·소실 객체만 LLM에 전달하여 토큰 효율과 분석 집중도를 "
        "동시에 확보한다. 시스템 프롬프트 제약부 [620]은 \"'DISAPPEARED' ≠ 파괴 확인\"이라는 의미론적 "
        "가드레일을 명시하여 LLM의 군사적 오해석을 억제한다. 8섹션 강제부 [630]은 표준 IMINT 형식을 "
        "강제해 분석관 친화 산출물을 보장한다. LLM 호출부 [710][720]은 동일 모델을 1차(영문) → 2차(한국어 "
        "번역)로 재호출하는 에이전트 패턴을 형성하며, 별도 번역 모델 없이 GPU·메모리 자원을 절감한다. "
        "메타 헤더 부착부 [730]은 모델명·세션ID·관측 시각 등 추적 가능한 메타데이터를 산출물 선두에 "
        "삽입해 감사 가능성을 확보한다.",
    )

    add_h3(doc, "도 5 (동작 순서도)")
    add_para(
        doc,
        "본 발명의 방법은 페어링 결과 입력(S100) → 그래프 누적(S200~S400) → 커뮤니티 탐지(S500) → "
        "이중 검색(S600~S800) → 프롬프트 조립(S900~S1000) → LLM 에이전트(S1100~S1200) → 최종 산출(S1300)의 "
        "6단계로 진행된다. S300의 카운터 누적은 LLM 비호출 결정론 단계, S500은 그래프 알고리즘 단계, "
        "S1100/S1200은 동일 LLM의 다단계 호출 단계로서, 본 발명은 이 세 종류의 이질적 연산을 단일 "
        "파이프라인 내에서 결합하는 구성을 그 진보적 특징으로 한다.",
    )

    # ───────── 권장 청구항 골격 ─────────
    add_h1(doc, "[참고] 권장 청구항 골격")
    add_table(
        doc,
        ["청구항", "핵심"],
        [
            ["독립항 1 (시스템)", "A + B + C + D + E 결합 시스템 — 본 발명의 가장 넓은 권리범위"],
            ["독립항 2 (방법)", "S100~S1300 순차 방법 청구"],
            ["독립항 3 (기록매체)", "컴퓨터 판독 가능 기록매체 청구"],
            ["종속항 a", "0.01° 격자 양자화"],
            ["종속항 b", "LLM-free 결정론 인덱싱 ★ (Microsoft GraphRAG 대비 진보성)"],
            ["종속항 c", "Louvain + co_occurred_with 가중치"],
            ["종속항 d", "Local + Global 이중 검색"],
            ["종속항 e", "~500 토큰 예산"],
            ["종속항 f", "new/disappeared 한정 LLM 전달"],
            ["종속항 g", "DISAPPEARED 의미 제약 ★ (도메인 안전성)"],
            ["종속항 h", "표준 IMINT 8섹션 강제"],
            ["종속항 i", "동일 LLM 번역 재호출"],
            ["종속항 j", "메타 헤더 부착으로 감사 가능성"],
        ],
        widths=[4.0, 12.0],
    )

    # ───────── 변리사 검증 추천 ─────────
    add_h1(doc, "[참고] 변리사 검증 추천 작업 (출원 전)")
    for s in [
        "KIPRIS 키워드 검색: 'LLM 보고서', 'RAG 보고서', '지식그래프 보고서', '위성영상 변화탐지 보고서', "
        "'IMINT 자동 생성'",
        "Google Patents 검색: 'intelligence report generation LLM', 'graph rag report', "
        "'change detection report automation'",
        "Microsoft GraphRAG 관련 실제 출원 여부 확인 (있다면 청구항 7로 회피)",
        "Palantir 정보 분석 자동화 특허군 검토",
        "SAM/SAM2/SAM3 관련 Meta 출원 검토",
    ]:
        add_bullet(doc, s)

    out_path = "/home/user/multi-source-intelligent-system/data/AI에이전트_판독보고서_특허명세서.docx"
    import os
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
